#!/usr/bin/env python3
"""Minimal Markdown → DOCX converter using python-docx.

Handles the subset of Markdown used in the SBA bonus docs:
  - # / ## / ### / #### headings
  - paragraphs
  - bulleted + numbered lists (- or 1.)
  - **bold**, *italic*
  - inline `code`
  - horizontal rules (---)
  - simple tables (| col | col |) — converted to docx tables

Usage:
    md_to_docx.py <input.md> <output.docx>
"""
from __future__ import annotations
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Brand colors
BRAND_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
GREY = RGBColor(0x80, 0x80, 0x80)


def _apply_inline(par, text):
    """Apply **bold**, *italic*, `code` formatting within a paragraph."""
    # Tokenize on ** | * | `
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = par.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = par.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = par.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            par.add_run(part)


def _add_table_from_md(doc, table_lines):
    """Convert a markdown table block to a docx table."""
    rows = []
    for line in table_lines:
        line = line.strip().strip("|")
        cells = [c.strip() for c in line.split("|")]
        rows.append(cells)
    if not rows:
        return
    # Drop separator row (---|---|---)
    if len(rows) >= 2 and all("-" in c or c == "" for c in rows[1]):
        rows.pop(1)
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = tbl.cell(i, j)
                cell.text = ""
                par = cell.paragraphs[0]
                if i == 0:
                    run = par.add_run(cell_text)
                    run.bold = True
                else:
                    _apply_inline(par, cell_text)


def convert(md_path, docx_path):
    md = Path(md_path).read_text()
    doc = Document()
    # Set base styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Smaller margins to maximize content
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    lines = md.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        # H1-H6
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = BRAND_BLUE
            if level == 1:
                run.font.size = Pt(22)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(14)
            elif level == 2:
                run.font.size = Pt(16)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
            elif level == 3:
                run.font.size = Pt(13)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
            else:
                run.font.size = Pt(11)
                p.paragraph_format.space_before = Pt(8)
            i += 1
            continue
        # Horizontal rule
        if re.match(r"^[\s]*[-*_]{3,}[\s]*$", line):
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(8)
            p_format.space_after = Pt(8)
            run = p.add_run("─" * 80)
            run.font.color.rgb = GREY
            i += 1
            continue
        # Table
        if line.strip().startswith("|") and line.count("|") >= 2:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|") and lines[i].count("|") >= 2:
                table_lines.append(lines[i])
                i += 1
            _add_table_from_md(doc, table_lines)
            continue
        # Bulleted list (- or *)
        if re.match(r"^[\s]*[-*]\s+", line):
            list_lines = []
            while i < len(lines) and re.match(r"^[\s]*[-*]\s+", lines[i]):
                content = re.sub(r"^[\s]*[-*]\s+", "", lines[i])
                list_lines.append(content)
                i += 1
            for item in list_lines:
                p = doc.add_paragraph(style="List Bullet")
                _apply_inline(p, item)
            continue
        # Numbered list (1.)
        if re.match(r"^[\s]*\d+\.\s+", line):
            list_lines = []
            while i < len(lines) and re.match(r"^[\s]*\d+\.\s+", lines[i]):
                content = re.sub(r"^[\s]*\d+\.\s+", "", lines[i])
                list_lines.append(content)
                i += 1
            for item in list_lines:
                p = doc.add_paragraph(style="List Number")
                _apply_inline(p, item)
            continue
        # Blank line — paragraph break already implied; skip
        if not line.strip():
            i += 1
            continue
        # Regular paragraph (may span multiple lines until blank)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(r"^(#{1,6}\s|[\s]*[-*]\s|[\s]*\d+\.\s|[\s]*\|)", lines[i]):
            para_lines.append(lines[i])
            i += 1
        text = " ".join(para_lines).strip()
        # Trailing double-space + newline = explicit break in markdown — collapse
        text = re.sub(r"\s+", " ", text)
        if text:
            p = doc.add_paragraph()
            _apply_inline(p, text)

    doc.save(docx_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("usage: md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
    print(f"OK — wrote {sys.argv[2]}")
