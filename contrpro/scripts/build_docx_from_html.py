#!/usr/bin/env python3
"""Batch convert ContrPro HTML legal + GC suite docs to Word .docx format.

Part of the integration-strategy execution locked 2026-05-17 — every doc
ships in Word + Google Docs + PDF (HTML stays as the source-of-truth). This
script handles the HTML→DOCX leg; Google Docs round-trip works natively
from the same .docx via "File → Open" or Drive upload.

Strategy:
  - Inputs: any .html under contrpro/files/packages/<tier>/documents/ or .../gc/
  - Output: <basename>.docx alongside the source HTML
  - Uses python-docx (already installed; no system pandoc required)
  - Preserves: H1-H6 headings, paragraphs, bold/italic, ordered + unordered
    lists, hyperlinks, basic table structure
  - Drops: HTML-specific CSS, inline styles, JS, images-from-data-URIs
  - Output is compatible with: Microsoft Word 2016+, Google Docs upload,
    Apple Pages, LibreOffice
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from bs4 import BeautifulSoup
except ImportError:
    print("ERROR: beautifulsoup4 required. Install: pip install beautifulsoup4", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx required. Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a hyperlink to a docx paragraph — python-docx doesn't expose this
    directly, so we manipulate the underlying XML."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def render_inline(paragraph, element) -> None:
    """Walk children of a block-level HTML element and emit inline runs.
    Handles <strong>/<b>, <em>/<i>, <a>, <br>, and bare text."""
    for child in element.children:
        if isinstance(child, str):
            paragraph.add_run(child)
            continue
        name = child.name
        if name in ("strong", "b"):
            paragraph.add_run(child.get_text()).bold = True
        elif name in ("em", "i"):
            paragraph.add_run(child.get_text()).italic = True
        elif name == "a":
            href = child.get("href", "").strip()
            text = child.get_text()
            if href and href.startswith(("http://", "https://", "mailto:")):
                add_hyperlink(paragraph, href, text)
            else:
                paragraph.add_run(text)
        elif name == "br":
            paragraph.add_run().add_break()
        elif name == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Courier New"
        else:
            # Unknown inline element — extract text content
            paragraph.add_run(child.get_text())


def convert(html_path: Path, docx_path: Path) -> dict:
    """Convert one HTML file to .docx. Returns stats."""
    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    doc = Document()

    # Set default font for body
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    body = soup.body or soup
    stats = {"headings": 0, "paragraphs": 0, "lists": 0, "tables": 0, "links": 0}

    def emit(element):
        name = getattr(element, "name", None)
        if name is None:
            return
        if name in ("script", "style", "head", "meta", "link"):
            return
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(name[1])
            heading = doc.add_heading(level=level)
            render_inline(heading, element)
            stats["headings"] += 1
        elif name == "p":
            p = doc.add_paragraph()
            render_inline(p, element)
            stats["paragraphs"] += 1
        elif name == "ul":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Bullet")
                render_inline(p, li)
            stats["lists"] += 1
        elif name == "ol":
            for li in element.find_all("li", recursive=False):
                p = doc.add_paragraph(style="List Number")
                render_inline(p, li)
            stats["lists"] += 1
        elif name == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            render_inline(p, element)
        elif name == "hr":
            doc.add_paragraph("―" * 30)
        elif name == "table":
            rows = element.find_all("tr")
            if not rows:
                return
            cols = max(len(r.find_all(["td", "th"])) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Light Grid Accent 1"
            for ri, tr in enumerate(rows):
                cells = tr.find_all(["td", "th"])
                for ci, cell in enumerate(cells):
                    if ri < len(table.rows) and ci < cols:
                        table.cell(ri, ci).text = cell.get_text(strip=True)
            stats["tables"] += 1
        elif name in ("div", "section", "article", "main", "header", "footer"):
            for child in element.children:
                emit(child)
        elif name == "pre":
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    for child in body.children:
        emit(child)

    # Count links across the whole document for stats
    stats["links"] = sum(1 for a in body.find_all("a") if a.get("href"))

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return stats


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", default="/Users/home/charles/contrpro/files/packages",
                    help="Root packages directory to walk for .html files")
    ap.add_argument("--dry-run", action="store_true", help="Just list what would be converted")
    args = ap.parse_args()

    root = Path(args.root)
    html_files = sorted([p for p in root.glob("**/*.html") if "documents" in p.parts or "gc" in p.parts])

    print(f"Found {len(html_files)} HTML files under {root}")
    print()

    seen_basenames = set()  # de-dupe across tiers (same doc lives in multiple tiers)
    converted = []
    for html_path in html_files:
        rel = html_path.relative_to(root)
        if args.dry_run:
            print(f"  would convert: {rel}")
            continue
        docx_path = html_path.with_suffix(".docx")
        try:
            stats = convert(html_path, docx_path)
            sz = docx_path.stat().st_size
            print(f"  ✓ {rel}  →  .docx ({sz//1024} KB, "
                  f"{stats['headings']}h, {stats['paragraphs']}p, "
                  f"{stats['lists']}lists, {stats['tables']}tbl, {stats['links']}links)")
            converted.append(docx_path)
        except Exception as e:
            print(f"  ✗ {rel}  FAILED: {type(e).__name__}: {e}")

    if not args.dry_run:
        print()
        print(f"Converted {len(converted)} files.")


if __name__ == "__main__":
    main()
